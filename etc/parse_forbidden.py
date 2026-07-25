# -*- coding: utf-8 -*-
"""noname_1.docx, noname_2.docx (불용문자 자료)를 파싱해 한자별 불용 사유를 병합한다.
두 파일 모두 실제 텍스트(이미지 아님)라 OCR이 필요 없다."""
import json
import re
import docx

HANJA_RE = re.compile(r"[㐀-䶿一-鿿豈-﫿]")


def parse_file1(path):
    """형식: <독음>(<한자>) : <사유>  (일부 줄은 이전 항목의 사유가 줄바꿈되어 이어짐)"""
    d = docx.Document(path)
    entries = {}
    last_key = None
    pat = re.compile(r"^(\S)\(([^)]+)\)\s*[:;]\s*(.+)$")
    lines = []
    for p in d.paragraphs:
        lines.extend(p.text.split("\n"))
    for text in lines:
        text = text.strip()
        if not text:
            continue
        m = pat.match(text)
        if m:
            reading, hanja, reason = m.groups()
            hanja_chars = HANJA_RE.findall(hanja)
            if not hanja_chars:
                continue
            hanja = hanja_chars[0]
            entries.setdefault(hanja, {"reading": reading, "reasons": []})
            entries[hanja]["reasons"].append(reason.strip())
            last_key = hanja
        else:
            # 이전 항목 사유의 연속 (예: "戌, 亥生은 더욱 凶하다.")
            if last_key is not None:
                entries[last_key]["reasons"].append(text)
    return entries


def parse_file2(path):
    """형식: <한자> <훈> <독음> : <사유>"""
    d = docx.Document(path)
    entries = {}
    for p in d.paragraphs:
        text = p.text.strip()
        if not text or ":" not in text:
            continue
        left, _, reason = text.partition(":")
        left = left.strip()
        toks = left.split()
        if len(toks) < 2:
            continue
        hanja_chars = HANJA_RE.findall(toks[0])
        if not hanja_chars:
            continue
        hanja = hanja_chars[0]
        reading = toks[-1]
        hun = " ".join(toks[1:-1])
        entries.setdefault(hanja, {"reading": reading, "hun": hun, "reasons": []})
        entries[hanja]["reasons"].append(reason.strip())
    return entries


def main():
    e1 = parse_file1("ETC/noname_1.docx")
    e2 = parse_file2("ETC/noname_2.docx")
    print(f"file1: {len(e1)}개 한자, file2: {len(e2)}개 한자")

    merged = {}
    for hanja, v in e1.items():
        merged[hanja] = {
            "char": hanja,
            "reading": v["reading"],
            "reasons": list(dict.fromkeys(v["reasons"])),  # 순서 유지 dedupe
        }
    for hanja, v in e2.items():
        if hanja in merged:
            for r in v["reasons"]:
                if r not in merged[hanja]["reasons"]:
                    merged[hanja]["reasons"].append(r)
        else:
            merged[hanja] = {
                "char": hanja,
                "reading": v["reading"],
                "reasons": list(dict.fromkeys(v["reasons"])),
            }

    print(f"병합 결과: {len(merged)}개 한자")
    with open("ETC/forbidden_hanja.json", "w", encoding="utf-8") as f:
        json.dump(list(merged.values()), f, ensure_ascii=False, indent=1)
    print("wrote ETC/forbidden_hanja.json")


if __name__ == "__main__":
    main()
